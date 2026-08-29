# -*- coding: utf-8 -*-
"""ISTANS 논문경진대회 국문 논문 .docx (한글에서 열어 .hwp로 저장 가능).
양식: 휴먼명조·본문 12pt, 줄간격 160%, 양쪽맞춤, 여백 위20·아래20·좌30·우30mm,
구성 표지-제목-논문요약-목차-본문-참고문헌-부록."""
import os, re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
FIGROOT = os.path.join(OUT, "..", "istans_work", "results", "public_ai_real")
BODY_FONT = "휴먼명조"
GOTHIC = "휴먼고딕"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)
sec.left_margin = Mm(30); sec.right_margin = Mm(30)

# default style
st = doc.styles["Normal"]
st.font.name = BODY_FONT; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
pf = st.paragraph_format
pf.line_spacing = 1.6; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def setfont(run, name, size, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, gothic=False, before=2, after=2, indent0=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.6
    for chunk, bold in _bold_split(text):
        r = p.add_run(chunk)
        setfont(r, GOTHIC if gothic else BODY_FONT, size, bold=bold or gothic)
    return p


def _bold_split(text):
    out = []
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part:
            out.append((part, i % 2 == 1))
    return out or [("", False)]


def heading(text, level):
    sizes = {3: 15, 4: 13, 5: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 3 else 6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    color = (18, 52, 90) if level == 3 else (17, 17, 17)
    setfont(r, GOTHIC, sizes[level], bold=True, color=color)
    if level == 3:
        pPr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "12345A"); pb.append(bottom); pPr.append(pb)
    return p


def table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph(); cp.paragraph_format.space_before = Pt(4)
        r = cp.add_run(caption); setfont(r, GOTHIC, 10.5, bold=True)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = c.add_run(h); setfont(rr, GOTHIC, 10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cp = cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for chunk, bold in _bold_split(str(v)):
                rr = cp.add_run(chunk); setfont(rr, BODY_FONT, 10.5, bold=bold)
                cp.paragraph_format.line_spacing = 1.3
    return t


def figure(name, caption, width_mm=120):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption); setfont(r, GOTHIC, 10.5, bold=True)
    path = os.path.join(FIGROOT, name)
    if os.path.exists(path):
        ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(path, width=Mm(width_mm))


def source(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); setfont(r, BODY_FONT, 10, color=(90, 90, 90))
    p.paragraph_format.line_spacing = 1.3


def pagebreak():
    doc.add_page_break()


def center(text, size, gothic=True, bold=True, before=0, after=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); setfont(r, GOTHIC if gothic else BODY_FONT, size, bold=bold)
    return p


TITLE = "산업통계로 본 공공 AI 정책의 경제 양극화 효과"
SUBTITLE = "한국 공식통계와 합성 인구를 이용한 이질적 에이전트 시뮬레이션 — ‘모두의 AI’ 사업에의 함의"

# ---------------- 표지 ----------------
for _ in range(6):
    doc.add_paragraph()
center("2026 ISTANS 논문경진대회", 13, before=0, after=40)
center(TITLE, 20, before=0, after=8)
center(SUBTITLE, 12, gothic=False, bold=False, after=40)
for _ in range(6):
    doc.add_paragraph()
center("2026.  8.", 12, gothic=False, bold=False, after=14)
center("소　속 :  삼성디스플레이", 12, gothic=False, bold=False, after=2)
center("성　명 :  최낙초", 12, gothic=False, bold=False)
pagebreak()

# ---------------- 제목 + 논문요약 ----------------
center(TITLE, 16, after=2)
center(SUBTITLE, 11, gothic=False, bold=False, after=6)
center("논문 요약", 14, after=4)
for t in [
 "인공지능(AI)의 급속한 확산은 산업별로 상이한 노출도와 보완성을 통해 고용과 소득에 비대칭적 영향을 준다. 2025년 기준 국내 생성형 AI 이용자는 2,300만 명을 넘어섰지만, 국민의 약 3분의 1은 여전히 AI를 사용하지 못하고 있으며, 이용자 상당수도 외산 무료 서비스에 의존한다. 이러한 배경에서 정부는 2026년 7월 전 국민에게 무료·무제한 국산 AI 챗봇과 공공서비스 AI 에이전트를 제공하는 ‘모두의 AI’ 프로젝트에 착수하였다. 그러나 같은 서비스를 제공받더라도 숙련, 직무 유연성, 프리미엄 모델 접근, AI 자본의 소유가 다르면 실효 이용과 경제적 성과는 벌어진다.",
 "본 연구의 목적은 **보편적 공공 AI 제공이 AI가 유발하는 경제 양극화를 완화할 수 있는지**를 산업통계에 근거해 정량적으로 규명하는 데 있다. 이를 위해 명목 접근과 실효 이용을 구분하는 이질적 에이전트 전이 모형을 구성하고, 모형의 핵심 구조 파라미터를 임의로 가정하지 않고 **공식 산업·경제 통계로 직접 보정**하였다. 통계청 KOSIS의 산업별 취업자와 산업·지역별 노동생산성, 한국은행 ECOS의 국민계정 노동소득분배율을 OpenAPI로 수집하여 산업 간 재배치 속도, 자본소득 비중, 산업·지역 생산성 격차를 추정하였다. 인구 구조는 공개 합성 페르소나를 익명 집계하여 표본으로 사용하되 실제 국민 표본으로 해석하지 않았다.",
 "시장 기준안과 다섯 개 공공 AI 정책조합을 동일 난수·동일 초기인구로 짝지어 비교하고, 40회 몬테카를로의 95% 신뢰구간·메커니즘 제거(ablation)·표본크기 스케일링으로 강건성을 검증하였다. 분석 결과, 산업통계로 보정한 한국의 산업 간 재배치 속도는 연 1.4%로 낮았고, 노동소득분배율은 2015년 62.3%에서 2024년 67.4%로 상승하였으며, 제조업 지역 간 생산성 격차는 상·하위 약 57%였다.",
 "이 실측 구조를 반영한 시뮬레이션에서 **재배치 속도가 낮음에도 모든 정책에서 양극화가 상승**하였다. 동인은 재배치의 속도가 아니라 산업·교육·자본에 내재한 격차 구조였다. **여섯 정책 중 종합안만이 양극화를 유의하게 완화**하였다(시장 대비 감소분 +0.0275, 95% 신뢰구간 [0.0252, 0.0298]로 0 배제). 메커니즘 분해는 **양극화 완화의 주력이 ‘AI 자본소유 환류’, 취약계층 고용 개선의 주력이 ‘교육·돌봄’**이라는 두 채널 구조를 드러냈다.",
 "본 연구는 산업통계를 정책 시뮬레이션의 파라미터로 직접 활용하여 정책 평가의 임의성을 줄이고 재현성을 높이는 방법을 제시하며, ‘모두의 AI’를 이용률이 아니라 분배·전환·권리 지표로 설계·평가해야 한다는 함의를 도출한다. 아울러 OpenAPI 확대·다차원 결합표 제공·메타데이터 표준화 등 ISTANS 활용성 제고 방안을 제언한다.",
]:
    para(t)
source("주제어: 산업통계, 공공 AI, 경제 양극화, 이질적 에이전트 시뮬레이션, 노동생산성, 노동소득분배율, 모두의 AI")
pagebreak()

# ---------------- 목차 ----------------
center("목　차", 13, after=4)
toc = [
 ("제1장 서론 및 연구 방법",0),("1. 연구의 배경과 목적",1),("2. 선행연구 검토",1),("3. 연구 방법과 분석 절차",1),
 ("제2장 산업통계 기반 파라미터 보정과 모형",0),("1. 활용한 산업통계와 파라미터 보정",1),("2. 이질적 에이전트 모형",1),
 ("제3장 시뮬레이션 분석 결과와 정책 함의",0),("1. 실인구 시뮬레이션 결과",1),("2. 정책 논의",1),("3. 한계와 결론",1),
 ("참고문헌",0),("부록: 재현성 명세 및 명제",0),
]
for text, lvl in toc:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(6*lvl); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); setfont(r, BODY_FONT, 11)
pagebreak()

# ---------------- 본문 ----------------
exec(open(os.path.join(OUT, "paper_body.py"), encoding="utf-8").read())

doc.save(os.path.join(OUT, "ISTANS_논문경진대회_논문_최낙초.docx"))
print("wrote docx")
